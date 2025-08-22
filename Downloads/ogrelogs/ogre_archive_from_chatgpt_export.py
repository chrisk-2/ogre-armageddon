#!/usr/bin/env python3
r"""
Ogre Archive Script (Self-Contained, robust)
- Auto-detects the latest ChatGPT export in Downloads (ZIP or folder).
- Or pass a path manually:
    python ogre_archive_from_chatgpt_export.py "C:\\path\\to\\chatgpt-export.zip"
- Generates Ogre Captain’s Log (Word doc) in OneDrive\ogrelogs.

Notes:
- Handles both classic "mapping" trees and future "messages" arrays.
- Safely extracts text from mixed content (strings, dict parts, code/tool messages).
"""

import os, sys, glob, zipfile, json, math
from pathlib import Path
from datetime import datetime
from typing import Any, Iterable, List, Dict, Optional
from docx import Document

# === CONFIG ===
ONEDRIVE_DIR = Path.home() / "OneDrive" / "ogrelogs"
OUT_NAME = "Ogre_Captains_Log_FULL.docx"

# Technical filter (set to False for ALL chats)
TECHNICAL_ONLY = True
INCLUDE_KEYWORDS = [
    "docker","ubuntu","debian","linux","onvif","rtsp","ffmpeg","python","flask",
    "websocket","socketio","camera","nvr","shinobi","frigate","armageddon",
    "starfleet","borg","section 31","jarvis","ollama","llama","ws-discovery"
]
EXCLUDE_KEYWORDS = ["lol","haha","joke"]

# === HELPERS ===
def find_latest_export() -> Optional[Path]:
    downloads = Path(os.path.expandvars(r"%USERPROFILE%")) / "Downloads"
    if not downloads.exists():
        downloads = Path.home() / "Downloads"

    candidates: List[Path] = []
    # Zips
    for pat in ("chatgpt-export*.zip", "chatgptexport*.zip", "chatgpt_export*.zip"):
        candidates.extend([Path(p) for p in glob.glob(str(downloads / pat))])
    # Folders
    for pat in ("chatgpt-export*", "chatgptexport*", "chatgpt_export*"):
        for p in glob.glob(str(downloads / pat)):
            pp = Path(p)
            if pp.is_dir():
                candidates.append(pp)

    if not candidates:
        return None
    candidates.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0]

def _safe_load_json_bytes(b: bytes) -> Any:
    try:
        return json.loads(b.decode("utf-8"))
    except UnicodeDecodeError:
        return json.loads(b.decode("utf-8", errors="replace"))

def _safe_read_text(p: Path) -> str:
    try:
        return p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return p.read_text(encoding="utf-8", errors="replace")

def load_conversations(export_path: Path) -> List[Dict[str, Any]]:
    """Return a list of conversation dicts."""
    if export_path.is_file() and export_path.suffix.lower() == ".zip":
        with zipfile.ZipFile(export_path, "r") as z:
            # Some exports place conversations.json at root
            name = "conversations.json"
            if name not in z.namelist():
                # fallbacks (rare)
                alt = [n for n in z.namelist() if n.endswith("/conversations.json")]
                if not alt:
                    print("[!] conversations.json not found in ZIP.")
                    return []
                name = alt[0]
            with z.open(name) as f:
                return json.load(f)
    elif export_path.is_dir():
        json_file = export_path / "conversations.json"
        if json_file.exists():
            return json.loads(_safe_read_text(json_file))
    else:
        print("[!] Invalid export path:", export_path)
    return []

def _coerce_timestamp(ts: Any) -> Optional[datetime]:
    """Handle float UNIX seconds, ms, or ISO strings; return datetime or None."""
    if ts is None:
        return None
    if isinstance(ts, (int, float)) and not math.isnan(float(ts)):
        # Heuristic: treat > 10^12 as ms
        if ts > 1_000_000_000_000:
            ts /= 1000.0
        try:
            return datetime.fromtimestamp(float(ts))
        except (OSError, OverflowError, ValueError):
            return None
    if isinstance(ts, str):
        # Try ISO parse (best effort)
        try:
            # Allow "2025-08-21T10:00:00Z" or similar
            from datetime import timezone
            dt = datetime.fromisoformat(ts.replace("Z","+00:00"))
            return dt
        except Exception:
            return None
    return None

def _join_lines(parts: Iterable[str]) -> str:
    return "\n".join(p for p in parts if p is not None and p != "")

def _extract_text_from_part(part: Any) -> Optional[str]:
    """Normalize any 'part' to text if possible."""
    if part is None:
        return None
    if isinstance(part, str):
        return part.strip()
    if isinstance(part, dict):
        # Common shapes:
        # {"content_type":"text","text":"..."} or {"type":"text","text":"..."}
        for key in ("text", "content", "title"):
            if key in part and isinstance(part[key], str):
                return part[key].strip()
        # Code-like:
        if "code" in part and isinstance(part["code"], str):
            return part["code"].strip()
        # Tool-calls or URL mentions:
        if "url" in part and isinstance(part["url"], str):
            return f"[link] {part['url'].strip()}"
        # Fall back to a readable subset
        # Avoid dumping huge JSON blobs; just a hint:
        kind = part.get("content_type") or part.get("type") or "object"
        return f"[{kind}]"
    # Other types → ignore
    return None

def extract_text_from_message(msg: Dict[str, Any]) -> Optional[str]:
    """Return cleaned text from a message node."""
    if not msg:
        return None
    content = msg.get("content") or {}
    # Newer format: content may have "parts" (list), or direct "text"
    if isinstance(content, dict):
        # Direct text path
        if "text" in content and isinstance(content["text"], str):
            return content["text"].strip()
        # Parts path
        parts = content.get("parts")
        if isinstance(parts, list):
            texts = []
            for p in parts:
                t = _extract_text_from_part(p)
                if t:
                    texts.append(t)
            if texts:
                return _join_lines(texts)
    # Fallback: if content itself is string
    if isinstance(content, str):
        return content.strip()
    return None

def passes_filter(text: str) -> bool:
    if not TECHNICAL_ONLY:
        return True
    lower = text.lower()
    if any(k in lower for k in INCLUDE_KEYWORDS):
        if not any(x in lower for x in EXCLUDE_KEYWORDS):
            return True
    return False

def _iter_messages_from_mapping(mapping: Dict[str, Any]) -> Iterable[Dict[str, Any]]:
    for node_id, node in mapping.items():
        msg = node.get("message")
        if msg:
            yield msg

def _iter_messages_generic(convo: Dict[str, Any]) -> Iterable[Dict[str, Any]]:
    """Yield message dicts whether the convo uses 'mapping' or 'messages'."""
    if "mapping" in convo and isinstance(convo["mapping"], dict):
        yield from _iter_messages_from_mapping(convo["mapping"])
    elif "messages" in convo and isinstance(convo["messages"], list):
        for msg in convo["messages"]:
            if isinstance(msg, dict):
                yield msg

def build_doc(conversations: List[Dict[str, Any]], out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Ogre Captain's Log", 0)

    # Sort convos by create_time if present
    def _sort_key(c: Dict[str, Any]):
        dt = _coerce_timestamp(c.get("create_time"))
        return dt or datetime.min
    conversations = sorted(conversations, key=_sort_key)

    for convo in conversations:
        title = convo.get("title") or "Untitled Conversation"
        dt = _coerce_timestamp(convo.get("create_time"))
        header = f"{title} ({dt.strftime('%Y-%m-%d %H:%M') if dt else 'Unknown'})"
        doc.add_heading(header, level=1)

        for msg in _iter_messages_generic(convo):
            role = (msg.get("author") or {}).get("role", "user").upper()
            text = extract_text_from_message(msg)
            if not text:
                continue
            if passes_filter(text):
                # Compact noisy whitespace
                cleaned = "\n".join(
                    line.rstrip() for line in text.splitlines()
                ).strip()
                # Cap pathological length per paragraph (optional)
                if cleaned:
                    doc.add_paragraph(f"[{role}] {cleaned}")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

def main():
    # Pick export path (arg or autodetect)
    if len(sys.argv) > 1:
        export_path = Path(sys.argv[1])
        if not export_path.exists():
            print("[!] Provided path does not exist:", export_path)
            sys.exit(2)
    else:
        export_path = find_latest_export()
        if not export_path:
            print("[!] Could not find ChatGPT export in Downloads.")
            sys.exit(1)
        print(f"[i] Auto-detected export: {export_path}")

    conversations = load_conversations(export_path)
    if not conversations:
        print("[!] No conversations found.")
        sys.exit(3)

    out_path = ONEDRIVE_DIR / OUT_NAME
    build_doc(conversations, out_path)
    print(f"[✓] Log created at: {out_path}")

if __name__ == "__main__":
    main()
